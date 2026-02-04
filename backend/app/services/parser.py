"""
TraceBridge AI - Document Parser Service
Handles PDF and DOCX text extraction with page-level metadata.
"""

import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class ParsedPage:
    """Represents a parsed page with text and metadata."""
    
    def __init__(self, text: str, page_number: Optional[int] = None):
        self.text = text
        self.page_number = page_number
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "page_number": self.page_number
        }


class ParsedDocument:
    """Represents a fully parsed document."""
    
    def __init__(self, filename: str, pages: List[ParsedPage]):
        self.filename = filename
        self.pages = pages
    
    @property
    def full_text(self) -> str:
        """Get concatenated text from all pages."""
        return "\n\n".join(page.text for page in self.pages if page.text.strip())
    
    @property
    def page_count(self) -> int:
        return len(self.pages)


def parse_pdf(file_path: str) -> ParsedDocument:
    """
    Parse a PDF file and extract text with page numbers.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        ParsedDocument with pages containing text and page numbers
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    pages = []
    
    try:
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            
            # Clean up the text
            text = text.strip()
            
            if text:  # Only add non-empty pages
                pages.append(ParsedPage(
                    text=text,
                    page_number=page_num + 1  # 1-indexed
                ))
        
        doc.close()
        
        logger.info(f"Parsed PDF '{path.name}': {len(pages)} pages with content")
        
    except Exception as e:
        logger.error(f"Error parsing PDF '{file_path}': {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    return ParsedDocument(filename=path.name, pages=pages)


def parse_docx(file_path: str) -> ParsedDocument:
    """
    Parse a DOCX file and extract text.
    Note: DOCX doesn't have native page numbers, so we treat the whole document as one page.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        ParsedDocument with extracted text
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    pages = []
    
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n\n".join(paragraphs)
        
        if full_text.strip():
            # DOCX doesn't have page numbers, so we use None
            pages.append(ParsedPage(
                text=full_text,
                page_number=None
            ))
        
        logger.info(f"Parsed DOCX '{path.name}': {len(paragraphs)} paragraphs")
        
    except Exception as e:
        logger.error(f"Error parsing DOCX '{file_path}': {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    return ParsedDocument(filename=path.name, pages=pages)


def parse_document(file_path: str) -> ParsedDocument:
    """
    Parse a document based on its file extension.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        ParsedDocument with extracted content
        
    Raises:
        ValueError: If file type is not supported
    """
    path = Path(file_path)
    extension = path.suffix.lower()
    
    if extension == ".pdf":
        return parse_pdf(file_path)
    elif extension in [".docx", ".doc"]:
        if extension == ".doc":
            logger.warning("Legacy .doc format detected. Attempting to parse as .docx")
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported types: .pdf, .docx")
